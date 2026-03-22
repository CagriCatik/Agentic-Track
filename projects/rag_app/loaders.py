"""Corpus loaders for heterogeneous source files."""

from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass
from html import unescape
from pathlib import Path
from typing import Any

SUPPORTED_SUFFIXES = frozenset(
    {
        ".pdf",
        ".txt",
        ".md",
        ".markdown",
        ".rst",
        ".html",
        ".htm",
        ".json",
        ".csv",
        ".docx",
        ".xlsx",
    }
)

_INVALID_AUTHOR_VALUES = {"", "unknown", "n/a", "none"}


@dataclass(frozen=True)
class SourceDocument:
    source_key: str
    source_name: str
    source_path: str
    source_type: str
    title: str
    author: str
    segments: list[str]


def discover_source_files(corpus_root: Path) -> list[Path]:
    return sorted(
        path
        for path in corpus_root.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES
    )


def source_key(corpus_root: Path, source_path: Path) -> str:
    return source_path.relative_to(corpus_root).as_posix()


def humanize_source_name(source_path: Path) -> str:
    return source_path.stem.replace("_", " ").replace("-", " ").strip() or source_path.name


def load_source_document(source_path: Path, *, corpus_root: Path) -> SourceDocument:
    suffix = source_path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(source_path, corpus_root=corpus_root)
    if suffix in {".txt", ".md", ".markdown", ".rst"}:
        return _load_text(source_path, corpus_root=corpus_root)
    if suffix in {".html", ".htm"}:
        return _load_html(source_path, corpus_root=corpus_root)
    if suffix == ".json":
        return _load_json(source_path, corpus_root=corpus_root)
    if suffix == ".csv":
        return _load_csv(source_path, corpus_root=corpus_root)
    if suffix == ".docx":
        return _load_docx(source_path, corpus_root=corpus_root)
    if suffix == ".xlsx":
        return _load_xlsx(source_path, corpus_root=corpus_root)
    raise ValueError(f"Unsupported source type: {source_path.suffix}")


def _load_pdf(source_path: Path, *, corpus_root: Path) -> SourceDocument:
    try:
        import fitz  # PyMuPDF
    except ModuleNotFoundError as exc:  # pragma: no cover
        raise RuntimeError(
            "pymupdf is required for PDF indexing. Install with: pip install pymupdf"
        ) from exc

    doc = fitz.open(str(source_path))
    metadata = doc.metadata or {}
    segments = [page.get_text("text") or "" for page in doc]
    body_text = "\n\n".join(segment for segment in segments if segment.strip())

    title = _sanitize_title(metadata.get("title"), source_path)
    if not title:
        title = guess_title(body_text, fallback=humanize_source_name(source_path))

    author = _sanitize_author(metadata.get("author"))
    if not author:
        author = guess_author(body_text)

    return SourceDocument(
        source_key=source_key(corpus_root, source_path),
        source_name=source_path.name,
        source_path=source_key(corpus_root, source_path),
        source_type="pdf",
        title=title,
        author=author,
        segments=segments or [""],
    )


def _load_text(source_path: Path, *, corpus_root: Path) -> SourceDocument:
    text = source_path.read_text(encoding="utf-8", errors="ignore")
    title = guess_title(text, fallback=humanize_source_name(source_path))
    author = guess_author(text)
    return SourceDocument(
        source_key=source_key(corpus_root, source_path),
        source_name=source_path.name,
        source_path=source_key(corpus_root, source_path),
        source_type=source_path.suffix.lower().lstrip("."),
        title=title,
        author=author,
        segments=[text],
    )


def _load_html(source_path: Path, *, corpus_root: Path) -> SourceDocument:
    raw_html = source_path.read_text(encoding="utf-8", errors="ignore")
    text = _strip_html(raw_html)
    title = _html_title(raw_html) or guess_title(text, fallback=humanize_source_name(source_path))
    author = guess_author(text)
    return SourceDocument(
        source_key=source_key(corpus_root, source_path),
        source_name=source_path.name,
        source_path=source_key(corpus_root, source_path),
        source_type="html",
        title=title,
        author=author,
        segments=[text],
    )


def _load_json(source_path: Path, *, corpus_root: Path) -> SourceDocument:
    payload = json.loads(source_path.read_text(encoding="utf-8", errors="ignore"))
    lines = _flatten_json(payload)
    text = "\n".join(lines)
    title = guess_title(text, fallback=humanize_source_name(source_path))
    return SourceDocument(
        source_key=source_key(corpus_root, source_path),
        source_name=source_path.name,
        source_path=source_key(corpus_root, source_path),
        source_type="json",
        title=title,
        author="",
        segments=[text],
    )


def _load_csv(source_path: Path, *, corpus_root: Path) -> SourceDocument:
    rows: list[str] = []
    with source_path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            items = [f"{key}: {value}" for key, value in row.items() if value not in {None, ""}]
            if items:
                rows.append(" | ".join(items))

    text = "\n".join(rows)
    title = humanize_source_name(source_path)
    return SourceDocument(
        source_key=source_key(corpus_root, source_path),
        source_name=source_path.name,
        source_path=source_key(corpus_root, source_path),
        source_type="csv",
        title=title,
        author="",
        segments=[text],
    )


def _load_docx(source_path: Path, *, corpus_root: Path) -> SourceDocument:
    try:
        from docx import Document as DocxDocument
    except ModuleNotFoundError as exc:  # pragma: no cover
        raise RuntimeError(
            "python-docx is required for DOCX indexing. Install with: pip install python-docx"
        ) from exc

    doc = DocxDocument(str(source_path))
    paragraphs = [" ".join(paragraph.text.split()).strip() for paragraph in doc.paragraphs]
    lines = [line for line in paragraphs if line]
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()).strip() for cell in row.cells]
            cleaned = [cell for cell in cells if cell]
            if cleaned:
                lines.append(" | ".join(cleaned))

    text = "\n".join(lines)
    title = _sanitize_title(doc.core_properties.title, source_path)
    if not title:
        title = guess_title(text, fallback=humanize_source_name(source_path))
    author = _sanitize_author(doc.core_properties.author)
    if not author:
        author = guess_author(text)

    return SourceDocument(
        source_key=source_key(corpus_root, source_path),
        source_name=source_path.name,
        source_path=source_key(corpus_root, source_path),
        source_type="docx",
        title=title,
        author=author,
        segments=_group_lines(lines) or [text],
    )


def _load_xlsx(source_path: Path, *, corpus_root: Path) -> SourceDocument:
    try:
        from openpyxl import load_workbook
    except ModuleNotFoundError as exc:  # pragma: no cover
        raise RuntimeError(
            "openpyxl is required for XLSX indexing. Install with: pip install openpyxl"
        ) from exc

    workbook = load_workbook(filename=str(source_path), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        header: list[str] | None = None
        for row in sheet.iter_rows(values_only=True):
            values = [_stringify_cell(value) for value in row]
            cleaned = [value for value in values if value]
            if not cleaned:
                continue
            if header is None and _looks_like_header_row(values):
                header = [value or f"column_{index + 1}" for index, value in enumerate(values)]
                continue

            if header:
                fields = [
                    f"{header[index]}: {value}"
                    for index, value in enumerate(values)
                    if value
                ]
            else:
                fields = [
                    f"column_{index + 1}: {value}"
                    for index, value in enumerate(values)
                    if value
                ]
            if fields:
                lines.append(f"Sheet {sheet.title} | " + " | ".join(fields))

    text = "\n".join(lines)
    title = _sanitize_title(workbook.properties.title, source_path)
    if not title:
        title = humanize_source_name(source_path)
    author = _sanitize_author(workbook.properties.creator)
    workbook.close()

    return SourceDocument(
        source_key=source_key(corpus_root, source_path),
        source_name=source_path.name,
        source_path=source_key(corpus_root, source_path),
        source_type="xlsx",
        title=title,
        author=author,
        segments=_group_lines(lines) or [text],
    )


def guess_title(text: str, *, fallback: str) -> str:
    lines = _candidate_lines(text)
    if not lines:
        return fallback

    first = lines[0]
    if (
        len(lines) > 1
        and not _looks_like_author_prefix(lines[1])
        and _looks_like_heading(lines[1])
        and len(first) + len(lines[1]) <= 140
    ):
        return f"{first} {lines[1]}"
    return first


def guess_author(text: str) -> str:
    for line in _candidate_lines(text, max_lines=12):
        for prefix in ("author:", "authors:", "by ", "von "):
            if _looks_like_author_prefix(line, prefix=prefix):
                candidate = line[len(prefix) :].strip(" :-")
                sanitized = _sanitize_author(candidate)
                if sanitized:
                    return sanitized
    return ""


def _candidate_lines(text: str, *, max_lines: int = 20) -> list[str]:
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = _normalize_text_line(raw_line)
        if not line:
            continue
        if len(line) > 160:
            continue
        if len(re.findall(r"[A-Za-z0-9ÄÖÜäöüß]", line)) < 3:
            continue
        lines.append(line)
        if len(lines) >= max_lines:
            break
    return lines


def _looks_like_heading(line: str) -> bool:
    words = re.findall(r"[A-Za-z0-9ÄÖÜäöüß'-]+", line)
    if not 2 <= len(words) <= 14:
        return False
    alpha_words = [word for word in words if any(ch.isalpha() for ch in word)]
    if not alpha_words:
        return False
    capitalized = sum(1 for word in alpha_words if word[0].isupper())
    return capitalized >= max(1, len(alpha_words) // 2)


def _looks_like_author_prefix(line: str, *, prefix: str | None = None) -> bool:
    lowered = _normalize_text_line(line).casefold()
    prefixes = (prefix,) if prefix is not None else ("author:", "authors:", "by ", "von ")
    return any(lowered.startswith(item) for item in prefixes)


def _normalize_text_line(raw_line: str) -> str:
    line = raw_line.strip()
    if not line:
        return ""
    line = re.sub(r"^\s{0,3}#{1,6}\s*", "", line)
    line = re.sub(r"^\s*[-*+]\s+", "", line)
    line = re.sub(r"^\s*\d+\.\s+", "", line)
    line = re.sub(r"\[(.*?)\]\([^)]+\)", r"\1", line)
    line = line.replace("**", "").replace("__", "").replace("`", "")
    line = " ".join(line.split()).strip(" -:_*")
    return line


def _sanitize_title(raw_title: Any, source_path: Path) -> str:
    if not raw_title:
        return ""
    title = " ".join(str(raw_title).split()).strip()
    if not title:
        return ""
    if title.casefold() == source_path.name.casefold():
        return ""
    return title


def _sanitize_author(raw_author: Any) -> str:
    if not raw_author:
        return ""
    author = " ".join(str(raw_author).split()).strip()
    if not author or author.casefold() in _INVALID_AUTHOR_VALUES:
        return ""
    if len(author) > 120:
        return ""
    return author


def _strip_html(raw_html: str) -> str:
    without_scripts = re.sub(
        r"<script.*?>.*?</script>|<style.*?>.*?</style>",
        " ",
        raw_html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    without_tags = re.sub(r"<[^>]+>", " ", without_scripts)
    normalized = unescape(without_tags)
    return re.sub(r"\s+", " ", normalized).strip()


def _html_title(raw_html: str) -> str:
    match = re.search(r"<title[^>]*>(.*?)</title>", raw_html, flags=re.IGNORECASE | re.DOTALL)
    if not match:
        return ""
    title = unescape(re.sub(r"\s+", " ", match.group(1))).strip()
    return title


def _flatten_json(value: Any, *, path: tuple[str, ...] = ()) -> list[str]:
    if isinstance(value, dict):
        lines: list[str] = []
        for key, nested in value.items():
            key_name = str(key)
            lines.extend(_flatten_json(nested, path=(*path, key_name)))
        return lines

    if isinstance(value, list):
        lines: list[str] = []
        for index, nested in enumerate(value):
            lines.extend(_flatten_json(nested, path=(*path, str(index))))
        return lines

    if value is None:
        return []

    key = ".".join(path)
    text = str(value).strip()
    if not text:
        return []
    if key:
        return [f"{key}: {text}"]
    return [text]


def _group_lines(lines: list[str], *, max_chars: int = 1800) -> list[str]:
    grouped: list[str] = []
    current: list[str] = []
    current_length = 0
    for line in lines:
        if not line:
            continue
        projected = current_length + len(line) + (1 if current else 0)
        if current and projected > max_chars:
            grouped.append("\n".join(current))
            current = [line]
            current_length = len(line)
            continue
        current.append(line)
        current_length = projected
    if current:
        grouped.append("\n".join(current))
    return grouped


def _stringify_cell(value: Any) -> str:
    if value is None:
        return ""
    return " ".join(str(value).split()).strip()


def _looks_like_header_row(values: list[str]) -> bool:
    populated = [value for value in values if value]
    if len(populated) < 2:
        return False
    if any(len(value) > 80 for value in populated):
        return False
    return all(not re.fullmatch(r"[0-9]+(?:\.[0-9]+)?", value) for value in populated)
