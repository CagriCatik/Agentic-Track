from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import Workbook

from projects.rag_app.loaders import discover_source_files, load_source_document


class LoaderTests(unittest.TestCase):
    def test_markdown_loader_infers_title_and_author_prefix(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            corpus_root = Path(temp_dir)
            source_path = corpus_root / "guide.md"
            source_path.write_text(
                "# Guide to Diagnostics\n\n- **Author: Jane Doe**\n\nThis guide explains diagnostic workflows.",
                encoding="utf-8",
            )

            loaded = load_source_document(source_path, corpus_root=corpus_root)

            self.assertEqual(loaded.title, "Guide to Diagnostics")
            self.assertEqual(loaded.author, "Jane Doe")
            self.assertEqual(loaded.source_type, "md")

    def test_json_loader_flattens_nested_content(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            corpus_root = Path(temp_dir)
            source_path = corpus_root / "service.json"
            source_path.write_text(
                json.dumps({"service": {"name": "catalog", "owner": "platform"}}),
                encoding="utf-8",
            )

            loaded = load_source_document(source_path, corpus_root=corpus_root)

            self.assertIn("service.name: catalog", loaded.segments[0])
            self.assertIn("service.owner: platform", loaded.segments[0])

    def test_docx_loader_reads_core_properties_and_table_content(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            corpus_root = Path(temp_dir)
            source_path = corpus_root / "runbook.docx"
            doc = DocxDocument()
            doc.core_properties.title = "Operations Runbook"
            doc.core_properties.author = "Jane Doe"
            doc.add_paragraph("This runbook explains incident response.")
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "System"
            table.rows[0].cells[1].text = "Owner"
            table.rows[1].cells[0].text = "Catalog"
            table.rows[1].cells[1].text = "Platform"
            doc.save(source_path)

            loaded = load_source_document(source_path, corpus_root=corpus_root)

            self.assertEqual(loaded.source_type, "docx")
            self.assertEqual(loaded.title, "Operations Runbook")
            self.assertEqual(loaded.author, "Jane Doe")
            self.assertIn("Catalog | Platform", "\n".join(loaded.segments))

    def test_xlsx_loader_reads_header_mapped_rows(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            corpus_root = Path(temp_dir)
            source_path = corpus_root / "inventory.xlsx"
            workbook = Workbook()
            workbook.properties.title = "Service Inventory"
            workbook.properties.creator = "Platform Team"
            sheet = workbook.active
            sheet.title = "Services"
            sheet.append(["service", "owner"])
            sheet.append(["catalog", "platform"])
            workbook.save(source_path)

            loaded = load_source_document(source_path, corpus_root=corpus_root)

            self.assertEqual(loaded.source_type, "xlsx")
            self.assertEqual(loaded.title, "Service Inventory")
            self.assertEqual(loaded.author, "Platform Team")
            self.assertIn("Sheet Services | service: catalog | owner: platform", "\n".join(loaded.segments))

    def test_discover_source_files_filters_supported_types(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            corpus_root = Path(temp_dir)
            (corpus_root / "a.md").write_text("hello", encoding="utf-8")
            (corpus_root / "b.txt").write_text("hello", encoding="utf-8")
            DocxDocument().save(corpus_root / "c.docx")
            (corpus_root / "ignore.bin").write_bytes(b"\x00\x01")

            discovered = [path.name for path in discover_source_files(corpus_root)]

            self.assertEqual(discovered, ["a.md", "b.txt", "c.docx"])


if __name__ == "__main__":
    unittest.main()
